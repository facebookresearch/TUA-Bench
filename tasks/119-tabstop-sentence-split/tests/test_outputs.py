# Copyright (c) Meta Platforms, Inc. and affiliates.
# All rights reserved.
#
# This source code is licensed under the license found in the
# LICENSE file in the root directory of this source tree.

TASK_SPEC = {
    "slug": "119-tabstop-sentence-split",
    "output_path": "/app/04 CHIN9505 EBook Purchasing info 2021 Jan.docx",
    "metric_spec": {
        "func": "check_tabstops",
        "result": "/app/04 CHIN9505 EBook Purchasing info 2021 Jan.docx",
        "options": {"word_number_split_by_tabstop": 3, "index": 0},
        "expected": {
            "type": "cloud_file",
            "url": "https://huggingface.co/datasets/xlangai/ubuntu_osworld_file_cache/resolve/main/libreoffice_writer/0a0faba3-5580-44df-965d-f562a99b291c/04%20CHIN9505%20EBook%20Purchasing%20info%202021%20Jan_Gold.docx",
            "dest": "04 CHIN9505 EBook Purchasing info 2021 Jan_Gold.docx",
        },
    },
}

import logging
import re
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT

logger = logging.getLogger("tua.writer.osworld")
logging.basicConfig(level=logging.INFO)

CACHE_DIR = Path("/tmp") / TASK_SPEC["slug"]


def fetch_cloud_file(url: str, dest: str) -> Path:
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    target = CACHE_DIR / dest
    if not target.exists():
        urllib.request.urlretrieve(url, target)
    return target


def check_tabstops(docx_file1, docx_file2, **kwargs):
    if not docx_file1 or not docx_file2:
        return 0.0

    try:
        doc1 = Document(docx_file1)
        doc2 = Document(docx_file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0.0

    para1 = [p for p in doc1.paragraphs if p.text.strip()]
    para2 = [p for p in doc2.paragraphs if p.text.strip()]
    if len(para1) != len(para2):
        return 0.0
    if not para1:
        return 0.0

    if kwargs.get("word_number_split_by_tabstop") is not None:
        number = kwargs["word_number_split_by_tabstop"]
        index = kwargs.get("index", 0)
        for p1 in para1:
            splits = p1.text.split("\t")
            if index >= len(splits):
                return 0.0
            words = [word for word in re.split(r"\s", splits[index]) if word.strip()]
            if len(words) != number:
                return 0.0

    section = doc2.sections[0]
    paragraph_width = section.page_width - section.left_margin - section.right_margin

    def ignore_tabs(tab_stop):
        return tab_stop.alignment == WD_TAB_ALIGNMENT.CLEAR or (
            tab_stop.alignment == WD_TAB_ALIGNMENT.LEFT and tab_stop.position == 0
        )

    minus = 0.0
    for p1, p2 in zip(para1, para2):
        tabs1 = [tab for tab in p1.paragraph_format.tab_stops if not ignore_tabs(tab)]
        tabs2 = [tab for tab in p2.paragraph_format.tab_stops if not ignore_tabs(tab)]
        if len(tabs1) != len(tabs2):
            return 0.0
        difference = 0.0
        for t1, t2 in zip(tabs1, tabs2):
            if t1.alignment != t2.alignment:
                return 0.0
            difference += abs(t1.position - t2.position)
        minus += difference / paragraph_width

    score = 1 - (minus / len(para1))
    return max(0.0, min(1.0, score))


def run_evaluation():
    metric_spec = TASK_SPEC["metric_spec"]
    expected_spec = metric_spec["expected"]
    expected_state = str(fetch_cloud_file(expected_spec["url"], expected_spec["dest"]))
    return float(
        check_tabstops(
            metric_spec["result"],
            expected_state,
            **metric_spec.get("options", {}),
        )
    )


if __name__ == "__main__":
    print(run_evaluation())
