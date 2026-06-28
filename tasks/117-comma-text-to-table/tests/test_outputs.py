# Copyright (c) Meta Platforms, Inc. and affiliates.
# All rights reserved.
#
# This source code is licensed under the license found in the
# LICENSE file in the root directory of this source tree.

TASK_SPEC = {
    "slug": "117-comma-text-to-table",
    "output_path": "/app/Graphemes_Sound_Letter_Patterns.docx",
    "metric_spec": {
        "func": ["compare_docx_tables", "compare_docx_tables"],
        "conj": "or",
        "result": [
            "/app/Graphemes_Sound_Letter_Patterns.docx",
            "/app/Graphemes_Sound_Letter_Patterns.docx",
        ],
        "expected": [
            {
                "type": "cloud_file",
                "url": "https://huggingface.co/datasets/xlangai/ubuntu_osworld_file_cache/resolve/main/libreoffice_writer/936321ce-5236-426a-9a20-e0e3c5dc536f/Graphemes_Sound_Letter_Patterns_Gold.docx",
                "dest": "Graphemes_Sound_Letter_Patterns_Gold.docx",
            },
            {
                "type": "cloud_file",
                "url": "https://huggingface.co/datasets/xlangai/ubuntu_osworld_file_cache/resolve/main/libreoffice_writer/936321ce-5236-426a-9a20-e0e3c5dc536f/Graphemes_Sound_Letter_Patterns_Gold_2.docx",
                "dest": "Graphemes_Sound_Letter_Patterns_Gold_2.docx",
            },
        ],
        "options": [{}, {}],
    },
}

import logging
import urllib.request
from pathlib import Path

from docx import Document

logger = logging.getLogger("tua.writer.osworld")
logging.basicConfig(level=logging.INFO)

CACHE_DIR = Path("/tmp") / TASK_SPEC["slug"]


def fetch_cloud_file(url: str, dest: str) -> Path:
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    target = CACHE_DIR / dest
    if not target.exists():
        urllib.request.urlretrieve(url, target)
    return target


def compare_docx_tables(docx_file1: str, docx_file2: str) -> float:
    if not docx_file1 or not docx_file2:
        return 0.0

    try:
        doc1 = Document(docx_file1)
        doc2 = Document(docx_file2)
    except Exception as exc:
        logger.error("Error: %s", exc)
        return 0.0

    tables1 = doc1.tables
    tables2 = doc2.tables
    if len(tables1) != len(tables2):
        return 0.0

    for table1, table2 in zip(tables1, tables2):
        if len(table1.rows) != len(table2.rows) or len(table1.columns) != len(table2.columns):
            return 0.0

        for row_idx in range(len(table1.rows)):
            for col_idx in range(len(table1.columns)):
                if table1.cell(row_idx, col_idx).text.strip() != table2.cell(row_idx, col_idx).text.strip():
                    return 0.0

    return 1.0


def resolve_expected(expected_spec: dict) -> str:
    if expected_spec["type"] != "cloud_file":
        raise ValueError(f"Unsupported expected spec: {expected_spec}")
    return str(fetch_cloud_file(expected_spec["url"], expected_spec["dest"]))


def evaluate_single(func_name: str, result_path: str, expected_spec: dict) -> float:
    if func_name != "compare_docx_tables":
        raise ValueError(f"Unsupported metric: {func_name}")
    return compare_docx_tables(result_path, resolve_expected(expected_spec))


def run_evaluation() -> float:
    metric_spec = TASK_SPEC["metric_spec"]
    funcs = metric_spec["func"]

    if isinstance(funcs, list):
        scores = [
            evaluate_single(func_name, metric_spec["result"][idx], metric_spec["expected"][idx])
            for idx, func_name in enumerate(funcs)
        ]
        conj = metric_spec.get("conj", "and")
        if conj == "or":
            return max(scores, default=0.0)
        if conj == "and":
            return min(scores, default=0.0)
        raise ValueError(f"Unsupported conjunction: {conj}")

    return evaluate_single(funcs, metric_spec["result"], metric_spec["expected"])


if __name__ == "__main__":
    print(run_evaluation())
