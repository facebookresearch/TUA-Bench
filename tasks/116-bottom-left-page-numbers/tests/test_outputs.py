# Copyright (c) Meta Platforms, Inc. and affiliates.
# All rights reserved.
#
# This source code is licensed under the license found in the
# LICENSE file in the root directory of this source tree.

TASK_SPEC = {'slug': '116-bottom-left-page-numbers',
 'output_path': '/app/LibreOffice_Open_Source_Word_Processing.docx',
 'metric_spec': {'func': 'has_page_numbers_in_footers',
                 'result': '/app/LibreOffice_Open_Source_Word_Processing.docx',
                 'options': {}}}

import logging
import re
import urllib.request
import xml.etree.ElementTree as ET
from io import BytesIO
from pathlib import Path
from typing import Any, List

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.shared import RGBColor

try:
    import fitz
except ImportError:
    fitz = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    from odf.opendocument import load
    from odf.text import P, Span
except ImportError:
    load = None
    P = None
    Span = None

try:
    from rapidfuzz import fuzz
except ImportError:
    fuzz = None

try:
    from skimage.color import deltaE_ciede2000
    from skimage.color import rgb2lab
except ImportError:
    deltaE_ciede2000 = None
    rgb2lab = None

logger = logging.getLogger("tua.writer.osworld")
logging.basicConfig(level=logging.INFO)


def _require_optional_dependencies(**modules: Any) -> None:
    missing = [name for name, module in modules.items() if module is None]
    if missing:
        raise RuntimeError(f"Missing optional verifier dependencies: {', '.join(missing)}")

LIMITATION_PHRASES = [
    "does not",
    "doesn't",
    "cannot",
    "can't",
    "not possible",
    "infeasible",
    "missing",
    "without",
    "unsupported",
    "unavailable",
    "not available",
    "not offered",
    "not listed",
    "nonexistent",
]

CACHE_DIR = Path("/tmp") / TASK_SPEC["slug"]


def fetch_cloud_file(url: str, dest: str) -> Path:
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    target = CACHE_DIR / dest
    if not target.exists():
        urllib.request.urlretrieve(url, target)
    return target


def find_default_font(config_file_path, rules):
    default_font = None
    expected_font = rules["font_name"]

    if not config_file_path:
        return 0

    try:
        tree = ET.parse(config_file_path)
        root = tree.getroot()
        namespace = {"oor": "http://openoffice.org/2001/registry"}

        for elem in root.findall(
            './/item[@oor:path="/org.openoffice.Office.Writer/DefaultFont"]',
            namespace,
        ):
            for prop in elem.findall('.//prop[@oor:name="Standard"]', namespace):
                for value in prop.findall("value", namespace):
                    default_font = value.text
    except Exception as e:
        logger.error(f"Error: {e}")

    return 1 if default_font == expected_font else 0


def contains_page_break(docx_file, rules):
    if not docx_file:
        return 0

    try:
        doc = Document(docx_file)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    expected_page_break_count = rules.get("page_break_count")
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    page_break_count = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            br_elems = run.element.findall(".//w:br", namespaces)
            for br in br_elems:
                if (
                    br is not None
                    and "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                    in br.attrib
                    and br.attrib[
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                    ]
                    == "page"
                ):
                    page_break_count += 1

    if expected_page_break_count is not None and page_break_count != expected_page_break_count:
        return 0

    return 1 if page_break_count > 0 else 0


def compare_docx_files(file1, file2, **options):
    ignore_blanks = options.get("ignore_blanks", True)
    ignore_case = options.get("ignore_case", False)
    ignore_order = options.get("ignore_order", False)
    content_only = options.get("content_only", False)
    fuzzy_match = options.get("fuzzy_match", False)
    delete_empty_lines = options.get("delete_empty_lines", False)

    if not file1 or not file2:
        return 0

    def get_paragraph_texts_odt(document):
        paragraphs = document.getElementsByType(P)
        paragraph_texts = []
        for paragraph in paragraphs:
            text_parts = []
            for node in paragraph.childNodes:
                if node.nodeType == node.TEXT_NODE:
                    text_parts.append(node.data)
                elif node.nodeType == node.ELEMENT_NODE and node.tagName == "text:span":
                    for child in node.childNodes:
                        if child.nodeType == child.TEXT_NODE:
                            text_parts.append(child.data)
            paragraph_texts.append("".join(text_parts))
        return paragraph_texts

    if str(file1).endswith(".docx") and str(file2).endswith(".docx"):
        try:
            doc1 = Document(file1)
            doc2 = Document(file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            return 0
        doc1_paragraphs = [p.text for p in doc1.paragraphs]
        doc2_paragraphs = [p.text for p in doc2.paragraphs]
        if ignore_order:
            doc1_paragraphs = sorted(doc1_paragraphs)
            doc2_paragraphs = sorted(doc2_paragraphs)
        if delete_empty_lines:
            doc1_paragraphs = [p for p in doc1_paragraphs if p.strip()]
            doc2_paragraphs = [p for p in doc2_paragraphs if p.strip()]
    elif str(file1).endswith(".odt") and str(file2).endswith(".odt"):
        _require_optional_dependencies(odfpy=load, odf_text=P)
        try:
            doc1 = load(file1)
            doc2 = load(file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            return 0
        doc1_paragraphs = get_paragraph_texts_odt(doc1)
        doc2_paragraphs = get_paragraph_texts_odt(doc2)
        if ignore_order:
            doc1_paragraphs = sorted(doc1_paragraphs)
            doc2_paragraphs = sorted(doc2_paragraphs)
        if delete_empty_lines:
            doc1_paragraphs = [p for p in doc1_paragraphs if p.strip()]
            doc2_paragraphs = [p for p in doc2_paragraphs if p.strip()]
    else:
        return 0

    if content_only:
        _require_optional_dependencies(rapidfuzz=fuzz)
        text1 = re.sub(r"\s+", " ", "\n".join(doc1_paragraphs)).strip()
        text2 = re.sub(r"\s+", " ", "\n".join(doc2_paragraphs)).strip()
        if ignore_case:
            text1, text2 = text1.lower(), text2.lower()
        return fuzz.ratio(text1, text2) / 100.0

    if ignore_blanks:
        text1 = re.sub(r"\s+", " ", "\n".join(doc1_paragraphs)).strip()
        text2 = re.sub(r"\s+", " ", "\n".join(doc2_paragraphs)).strip()
        if ignore_case:
            text1, text2 = text1.lower(), text2.lower()
        if fuzzy_match:
            _require_optional_dependencies(rapidfuzz=fuzz)
            return fuzz.ratio(text1, text2) / 100.0
        return 1 if text1 == text2 else 0

    if len(doc1_paragraphs) != len(doc2_paragraphs):
        return 0

    if fuzzy_match:
        _require_optional_dependencies(rapidfuzz=fuzz)
        if not doc1_paragraphs:
            return 1.0
        total_similarity = 0.0
        for p1, p2 in zip(doc1_paragraphs, doc2_paragraphs):
            if ignore_case:
                p1, p2 = p1.lower(), p2.lower()
            total_similarity += fuzz.ratio(p1, p2) / 100.0
        return total_similarity / len(doc1_paragraphs)

    for p1, p2 in zip(doc1_paragraphs, doc2_paragraphs):
        if ignore_case:
            p1, p2 = p1.lower(), p2.lower()
        if p1 != p2:
            return 0

    return 1


def compare_docx_tables(docx_file1, docx_file2):
    if not docx_file1 or not docx_file2:
        return 0

    try:
        doc1 = Document(docx_file1)
        doc2 = Document(docx_file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    tables1 = doc1.tables
    tables2 = doc2.tables

    if len(tables1) != len(tables2):
        return 0

    for table1, table2 in zip(tables1, tables2):
        if len(table1.rows) != len(table2.rows) or len(table1.columns) != len(table2.columns):
            return 0

        for i in range(len(table1.rows)):
            for j in range(len(table1.columns)):
                if table1.cell(i, j).text.strip() != table2.cell(i, j).text.strip():
                    return 0

    return 1


def compare_docx_images(docx_file1, docx_file2):
    if not docx_file1 or not docx_file2:
        return 0

    try:
        doc1 = Document(docx_file1)
        doc2 = Document(docx_file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    def extract_images(doc):
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_data = rel.target_part.blob
                images.append(BytesIO(img_data))
        return images

    images1 = extract_images(doc1)
    images2 = extract_images(doc2)
    if len(images1) != len(images2):
        return 0

    _require_optional_dependencies(pillow=Image)
    for img1, img2 in zip(images1, images2):
        if Image.open(img1).tobytes() != Image.open(img2).tobytes():
            return 0

    return 1


def compare_line_spacing(docx_file1, docx_file2):
    if not docx_file1 or not docx_file2:
        return 0

    if not compare_docx_files(docx_file1, docx_file2):
        return 0

    try:
        doc1 = Document(docx_file1)
        doc2 = Document(docx_file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    if len(doc1.paragraphs) != len(doc2.paragraphs):
        return 0

    for para1, para2 in zip(doc1.paragraphs, doc2.paragraphs):
        if para1.paragraph_format.line_spacing != para2.paragraph_format.line_spacing:
            return 0

    return 1


def compare_font_names(docx_file, rules):
    if not docx_file:
        return 0

    try:
        doc = Document(docx_file)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    expected_font = rules["font_name"]

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name != expected_font:
                return 0

    return 1


def compare_subscript_contains(docx_file1, docx_file2):
    if not docx_file1 or not docx_file2:
        return 0

    try:
        doc1 = Document(docx_file1)
        doc2 = Document(docx_file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    for para1, para2 in zip(doc1.paragraphs, doc2.paragraphs):
        for run1, run2 in zip(para1.runs, para2.runs):
            if run1.font.subscript and run2.font.subscript:
                return 1

    return 0


def has_page_numbers_in_footers(docx_file):
    if not docx_file:
        return 0

    try:
        doc = Document(docx_file)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    for section in doc.sections:
        footer = section.footer
        if footer is None:
            return 0
        footer_text = footer.paragraphs[0].text if footer.paragraphs else ""
        if not any(char.isdigit() for char in footer_text):
            return 0

    return 1


def is_first_line_centered(docx_file):
    if not docx_file:
        return 0

    try:
        doc = Document(docx_file)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    return 1 if doc.paragraphs[0].paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER else 0


def check_tabstops(docx_file1, docx_file2, **kwargs):
    if not docx_file1 or not docx_file2:
        return 0.0

    try:
        doc1 = Document(docx_file1)
        doc2 = Document(docx_file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0.0

    para1 = [p for p in doc1.paragraphs if p.text.strip()]
    para2 = [p for p in doc2.paragraphs if p.text.strip()]
    if len(para1) != len(para2):
        return 0.0

    if kwargs.get("word_number_split_by_tabstop") is not None:
        number = kwargs["word_number_split_by_tabstop"]
        index = kwargs.get("index", 0)
        for p1 in para1:
            splits = p1.text.split("\t")
            if len(splits) == 0:
                return 0.0
            words = [word for word in re.split(r"\s", splits[index]) if word.strip()]
            if len(words) != number:
                return 0.0

    section = doc2.sections[0]
    paragraph_width = section.page_width - section.left_margin - section.right_margin

    def ignore_tabs(tab_stop):
        return tab_stop.alignment == WD_TAB_ALIGNMENT.CLEAR or (
            tab_stop.alignment == WD_TAB_ALIGNMENT.LEFT and tab_stop.position == 0
        )

    minus = 0.0
    for p1, p2 in zip(para1, para2):
        tabs1 = [tab for tab in p1.paragraph_format.tab_stops if not ignore_tabs(tab)]
        tabs2 = [tab for tab in p2.paragraph_format.tab_stops if not ignore_tabs(tab)]
        if len(tabs1) != len(tabs2):
            return 0.0
        difference = 0.0
        for t1, t2 in zip(tabs1, tabs2):
            if t1.alignment != t2.alignment:
                return 0.0
            difference += abs(t1.position - t2.position)
        minus += difference / paragraph_width

    return 1 - (minus / len(para1))


def evaluate_colored_words_in_tables(file_path1, file_path2, **kwargs):
    if not file_path1 or not file_path2:
        return 0

    if not compare_docx_files(file_path1, file_path2):
        return 0

    _require_optional_dependencies(scikit_image_deltaE=deltaE_ciede2000, scikit_image_rgb2lab=rgb2lab)
    try:
        document = Document(file_path1)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    threshold = kwargs.get("threshold", 3.5)

    def calculate_color_difference(rgb1, rgb2):
        srgb1 = [rgb1[0] / 255.0, rgb1[1] / 255.0, rgb1[2] / 255.0]
        srgb2 = [rgb2[0] / 255.0, rgb2[1] / 255.0, rgb2[2] / 255.0]
        lab1, lab2 = rgb2lab(srgb1), rgb2lab(srgb2)
        return deltaE_ciede2000(lab1, lab2)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        word = run.text
                        if not word:
                            continue
                        first_letter = word[0].lower()
                        if first_letter in "aeiou":
                            if calculate_color_difference(run.font.color.rgb, RGBColor(255, 0, 0)) > threshold:
                                return 0
                        else:
                            if calculate_color_difference(run.font.color.rgb, RGBColor(0, 0, 255)) > threshold:
                                return 0

    return 1


def check_highlighted_words(file_path1, file_path2):
    if not file_path1 or not file_path2:
        return 0

    if not compare_docx_files(file_path1, file_path2):
        return 0

    _require_optional_dependencies(odfpy=load, odf_text=Span)
    doc = load(file_path1)
    highlighted = False

    for span in doc.getElementsByType(Span):
        style_name = span.getAttribute("stylename")
        if not style_name:
            continue
        for automatic_style in doc.automaticstyles.childNodes:
            if automatic_style.getAttribute("name") == style_name:
                for prop in automatic_style.childNodes:
                    if prop.getAttribute("backgroundcolor") == "#ffff00":
                        highlighted = True
                        break
        if highlighted:
            break

    return 0 if highlighted else 1


def evaluate_strike_through_last_paragraph(file_path1, file_path2):
    if not file_path1 or not file_path2:
        return 0

    if not compare_docx_files(file_path1, file_path2):
        return 0

    try:
        document = Document(file_path1)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    last_paragraph = document.paragraphs[-1]
    for run in last_paragraph.runs:
        if not run.font.strike:
            return 0

    return 1


def check_italic_font_size_14(path1, path2):
    if not path1 or not path2:
        return 0

    if not compare_docx_files(path1, path2):
        return 0

    try:
        document = Document(path1)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.italic:
                if run.font.size is None or run.font.size.pt != 14:
                    return 0

    return 1


def compare_unique_train_records(processed_file, expected_files, **kwargs):
    logger.info(f"DEBUG: processed_file type: {type(processed_file)}, value: {processed_file}")
    logger.info(f"DEBUG: expected_files type: {type(expected_files)}, value: {expected_files}")
    logger.info(f"DEBUG: kwargs: {kwargs}")

    if not processed_file or not isinstance(expected_files, list) or len(expected_files) < 2:
        logger.error("Invalid arguments: processed_file and a list of 2 expected_files are required.")
        return 0

    gold_file = expected_files[0]
    initial_file = expected_files[1]
    if not gold_file or not initial_file:
        logger.error("Gold file or initial file path is missing from expected_files list.")
        return 0

    def get_lines_and_ids(file_path):
        try:
            doc = Document(file_path)
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            train_ids = [line.split(",")[1].strip() for line in lines if len(line.split(",")) == 4]
            return lines, train_ids
        except Exception as e:
            logger.error(f"Error opening or parsing file {file_path}: {e}")
            return None, None

    processed_lines, processed_train_ids = get_lines_and_ids(processed_file)
    if processed_lines is None:
        return 0

    gold_lines, gold_train_ids = get_lines_and_ids(gold_file)
    if gold_lines is None:
        return 0

    initial_lines, _ = get_lines_and_ids(initial_file)
    if initial_lines is None:
        return 0

    if not set(processed_lines).issubset(set(initial_lines)):
        return 0
    if len(processed_train_ids) != len(set(processed_train_ids)):
        return 0
    if set(processed_train_ids) != set(gold_train_ids):
        return 0

    return 1


def compare_pdfs(pdf1_path, pdf2_path):
    if type(pdf2_path) != list:
        pdf1_path, pdf2_path = [pdf1_path], [pdf2_path]

    _require_optional_dependencies(pymupdf=fitz, rapidfuzz=fuzz)

    def extract_text_from_pdf(pdf_path):
        text = ""
        with fitz.open(pdf_path) as pdf:
            for page in pdf:
                text += page.get_text()
        return text.strip()

    score = 0.0
    for path1, path2 in zip(pdf1_path, pdf2_path):
        try:
            text1 = extract_text_from_pdf(path1)
            text2 = extract_text_from_pdf(path2)
            score += fuzz.ratio(text1, text2) / 100
        except Exception as e:
            logger.info(f"[ERROR]: unexpected error occurred when comparing PDF files: {e}")

    return score / len(pdf2_path)


def check_infeasible_response(path: Path, required_terms: List[str]) -> float:
    if path is None or not path.exists():
        return 0.0

    text = path.read_text(encoding="utf-8").strip()
    if not text:
        return 0.0

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines or lines[0].lower() != "infeasible":
        return 0.0

    reason = " ".join(lines[1:]).lower()
    if not reason:
        return 0.0
    if not all(term in reason for term in required_terms):
        return 0.0
    if not any(phrase in reason for phrase in LIMITATION_PHRASES):
        return 0.0

    return 1.0


def resolve_expected(spec):
    if spec is None:
        return None
    if spec["type"] == "cloud_file":
        return str(fetch_cloud_file(spec["url"], spec["dest"]))
    if spec["type"] == "cloud_file_list":
        return [
            str(fetch_cloud_file(url, dest))
            for url, dest in zip(spec["urls"], spec["dests"])
        ]
    if spec["type"] == "rule":
        return spec["rules"]
    raise ValueError(f"Unsupported expected spec: {spec}")


def evaluate_single(func_name, result_state, expected_state=None, options=None):
    options = options or {}
    metric = {
        "compare_docx_files": compare_docx_files,
        "compare_subscript_contains": compare_subscript_contains,
        "compare_line_spacing": compare_line_spacing,
        "check_tabstops": check_tabstops,
        "has_page_numbers_in_footers": has_page_numbers_in_footers,
        "compare_font_names": compare_font_names,
        "is_first_line_centered": is_first_line_centered,
        "compare_pdfs": compare_pdfs,
        "compare_docx_tables": compare_docx_tables,
        "check_highlighted_words": check_highlighted_words,
        "compare_docx_images": compare_docx_images,
        "compare_unique_train_records": compare_unique_train_records,
        "evaluate_strike_through_last_paragraph": evaluate_strike_through_last_paragraph,
        "evaluate_colored_words_in_tables": evaluate_colored_words_in_tables,
        "check_italic_font_size_14": check_italic_font_size_14,
        "contains_page_break": contains_page_break,
        "find_default_font": find_default_font,
    }[func_name]

    if expected_state is None:
        return float(metric(result_state, **options))
    return float(metric(result_state, expected_state, **options))


def run_evaluation():
    metric_spec = TASK_SPEC["metric_spec"]
    if metric_spec.get("mode") == "infeasible":
        return check_infeasible_response(Path(TASK_SPEC["output_path"]), metric_spec["required_terms"])

    func = metric_spec["func"]
    if isinstance(func, list):
        conj = metric_spec.get("conj", "and")
        results = []
        for idx, func_name in enumerate(func):
            result_state = metric_spec["result"][idx]
            expected_state = resolve_expected(metric_spec["expected"][idx]) if "expected" in metric_spec else None
            options = metric_spec.get("options", [{}] * len(func))[idx]
            metric_value = evaluate_single(func_name, result_state, expected_state, options)
            if conj == "and" and metric_value == 0.0:
                return 0.0
            if conj == "or" and metric_value == 1.0:
                return 1.0
            results.append(metric_value)
        return sum(results) / len(results) if conj == "and" else max(results)

    result_state = metric_spec["result"]
    expected_state = resolve_expected(metric_spec["expected"]) if "expected" in metric_spec else None
    return evaluate_single(func, result_state, expected_state, metric_spec.get("options", {}))


if __name__ == "__main__":
    print(run_evaluation())
